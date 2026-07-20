"""Human-readable Word (.docx) report of a learned context.

The ``.epwcontext`` package is a ZIP for re-import/sharing; this is the openable
counterpart — a formatted document of the application's cubes, dimensions,
member hierarchies, rules and variables that opens in Word / Pages / Google Docs.
"""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from sqlalchemy.orm import Session

from ..db.models import ContextVersion
from ..services import context_store

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# Color palette matching the app's IBM Carbon design system
COLOR_PRIMARY = RGBColor(69, 137, 255)      # #4589ff
COLOR_SUCCESS = RGBColor(66, 190, 101)      # #42be65
COLOR_WARNING = RGBColor(255, 131, 43)      # #ff832b
COLOR_TEXT_PRIMARY = RGBColor(22, 22, 22)   # #161616
COLOR_TEXT_SECONDARY = RGBColor(141, 141, 141)  # #8d8d8d
COLOR_BORDER = RGBColor(57, 57, 57)         # #393939


def _sorted_records(records: list, kind: str) -> list:
    return [r for r in records if r.kind == kind]


def _set_cell_shading(cell, color: RGBColor):
    """Add shading/background color to a table cell."""
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), f"{color.red:02X}{color.green:02X}{color.blue:02X}")
    cell._element.get_or_add_tcPr().append(shading_elm)


def _add_kv(doc, label: str, value: str, color: RGBColor | None = None) -> None:
    """Add a key-value pair with optional color."""
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    if color:
        run_label.font.color.rgb = color
    run_value = p.add_run(value)
    run_value.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)


def _add_table(doc, headers: list[str], rows: list[list[str]], header_color: RGBColor | None = None):
    """Add a formatted table with styled headers."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"

    # Style header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        if header_color:
            _set_cell_shading(cell, header_color)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255) if header_color else COLOR_TEXT_PRIMARY

    # Add data rows
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for para in cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
            # Alternate row shading for readability
            if row_idx % 2 == 0:
                _set_cell_shading(cells[i], RGBColor(249, 249, 249))

    return table


def _add_cube_diagram(doc, cubes_data: list, dims_data: list) -> None:
    """Add a visual diagram showing cube-dimension relationships."""
    # Build dimension -> cubes mapping
    dim_to_cubes: dict[str, list[str]] = {}
    for cube in cubes_data:
        data = cube.data or {}
        dims = data.get("dimensions") or []
        for dim in dims:
            dim_to_cubes.setdefault(dim, []).append(cube.name)

    # Create a text-based visualization
    doc.add_heading("Architecture Diagram", level=2)
    doc.add_paragraph("Visual representation of how dimensions connect to cubes:", style="Body Text")

    # Use a monospace paragraph for the diagram
    for cube in cubes_data[:10]:  # Limit to first 10 cubes for readability
        data = cube.data or {}
        dims = data.get("dimensions") or []
        cube_type = data.get("type", "")

        # Cube header
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"┌─ {cube.name}")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY

        # Cube type
        if cube_type:
            p_type = doc.add_paragraph()
            p_type.paragraph_format.space_after = Pt(2)
            run_type = p_type.add_run(f"│  Type: {cube_type}")
            run_type.font.name = "Courier New"
            run_type.font.size = Pt(8)
            run_type.font.color.rgb = COLOR_TEXT_SECONDARY

        # Dimensions
        for i, dim in enumerate(dims[:15]):  # Limit dimensions shown
            p_dim = doc.add_paragraph()
            p_dim.paragraph_format.space_after = Pt(1)
            prefix = "└─" if i == len(dims) - 1 or i == 14 else "├─"
            run_dim = p_dim.add_run(f"│  {prefix} {dim}")
            run_dim.font.name = "Courier New"
            run_dim.font.size = Pt(8)
            run_dim.font.color.rgb = COLOR_TEXT_PRIMARY

        if len(dims) > 15:
            p_more = doc.add_paragraph()
            p_more.paragraph_format.space_after = Pt(2)
            run_more = p_more.add_run(f"│     ... and {len(dims) - 15} more")
            run_more.font.name = "Courier New"
            run_more.font.size = Pt(8)
            run_more.font.color.rgb = COLOR_TEXT_SECONDARY
            run_more.italic = True

    doc.add_paragraph()  # Spacing


def _member_outline(doc, members: list, max_depth: int = 5) -> None:
    """Render one dimension's members as an indented hierarchy with tree characters."""
    by_name = {m.name: m for m in members}
    children: dict[str | None, list] = {}
    for m in members:
        parent = m.parent if (m.parent and m.parent in by_name and m.parent != m.name) else None
        children.setdefault(parent, []).append(m)
    roots = children.get(None, [])

    def walk(node, depth: int, is_last: bool = False, prefix: str = "") -> None:
        if depth > max_depth:
            return

        data = node.data or {}
        label = node.name
        alias = node.alias or data.get("alias")
        if alias and alias != node.name:
            label += f"  ({alias})"

        # Use tree characters for visual hierarchy
        connector = "└─ " if is_last else "├─ "
        full_prefix = prefix + connector

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.1)

        run = p.add_run(full_prefix + label)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        if depth == 0:
            run.font.bold = True
            run.font.color.rgb = COLOR_PRIMARY
        else:
            run.font.color.rgb = COLOR_TEXT_PRIMARY

        # Prepare prefix for children
        child_prefix = prefix + ("   " if is_last else "│  ")
        child_nodes = children.get(node.name, [])

        for idx, child in enumerate(child_nodes):
            walk(child, depth + 1, idx == len(child_nodes) - 1, child_prefix)

    for idx, root in enumerate(roots):
        walk(root, 0, idx == len(roots) - 1)


def build_context_docx(session: Session, context_version_id: str) -> tuple[str, bytes]:
    cv = session.get(ContextVersion, context_version_id)
    if cv is None:
        raise KeyError("context version not found")
    records = context_store.get_records(session, context_version_id)
    counts = (cv.manifest or {}).get("counts", {})

    doc = Document()

    # Title page with styled heading
    title = doc.add_heading(f"EPM Context Report", level=0)
    for run in title.runs:
        run.font.color.rgb = COLOR_PRIMARY
        run.font.size = Pt(28)

    subtitle = doc.add_paragraph(cv.application)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = COLOR_TEXT_SECONDARY

    doc.add_paragraph()  # Spacing

    # Metadata with colored labels
    _add_kv(doc, "Application", cv.application, COLOR_PRIMARY)
    _add_kv(doc, "Environment", (cv.manifest or {}).get("environmentClassification", "—"), COLOR_SUCCESS)
    _add_kv(doc, "Generated", (cv.manifest or {}).get("generatedAt", str(cv.created_at)))
    _add_kv(doc, "Context version", cv.label)

    doc.add_page_break()

    # --- summary with visual stats ---
    doc.add_heading("Summary", level=1)
    summary_para = doc.add_paragraph()
    summary_para.add_run("Overview of the EPM context contents.").font.size = Pt(10)

    _add_table(
        doc,
        ["Category", "Count"],
        [[k.replace("_", " ").title(), str(v)] for k, v in counts.items()],
        header_color=COLOR_PRIMARY,
    )

    doc.add_page_break()

    # --- cubes with diagram ---
    cubes = _sorted_records(records, "cube")
    if cubes:
        doc.add_heading(f"Cubes ({len(cubes)})", level=1)

        # Add architecture diagram
        dims = _sorted_records(records, "dimension")
        if cubes and dims:
            _add_cube_diagram(doc, cubes, dims)

        # Summary table
        doc.add_heading("Cube Details", level=2)
        rows = []
        for c in cubes:
            data = c.data or {}
            dims = data.get("dimensions") or []
            dim_list = ", ".join(dims[:5])
            if len(dims) > 5:
                dim_list += f" ... +{len(dims) - 5} more"
            rows.append([c.name, data.get("type", ""), str(len(dims)), dim_list])
        _add_table(
            doc,
            ["Cube", "Type", "# Dimensions", "Dimensions (first 5)"],
            rows,
            header_color=COLOR_PRIMARY,
        )

    doc.add_page_break()

    # --- dimensions ---
    dims = _sorted_records(records, "dimension")
    if dims:
        doc.add_heading(f"Dimensions ({len(dims)})", level=1)
        rows = []
        for d in dims:
            data = d.data or {}
            cube_list = ", ".join((data.get("cubes") or [])[:3])
            if len(data.get("cubes") or []) > 3:
                cube_list += f" +{len(data.get('cubes', [])) - 3}"
            rows.append(
                [
                    d.name,
                    data.get("type", ""),
                    "Dense" if data.get("dense") else "Sparse",
                    cube_list,
                ]
            )
        _add_table(
            doc,
            ["Dimension", "Type", "Density", "Used in cubes"],
            rows,
            header_color=COLOR_PRIMARY,
        )

    doc.add_page_break()

    # --- members, grouped by dimension as visual hierarchy trees ---
    members = _sorted_records(records, "member")
    if members:
        doc.add_heading(f"Member Hierarchies ({len(members)} total)", level=1)

        intro = doc.add_paragraph()
        intro.add_run(
            "Visual tree representation of dimension members showing parent-child relationships. "
            "Hierarchies are limited to 5 levels for readability."
        ).font.size = Pt(10)

        by_dim: dict[str, list] = {}
        for m in members:
            by_dim.setdefault(m.dimension or "—", []).append(m)

        # Show first 10 dimensions with hierarchies
        shown_dims = list(sorted(by_dim.keys()))[:10]

        for dim in shown_dims:
            doc.add_heading(f"{dim} ({len(by_dim[dim])} members)", level=2)
            _member_outline(doc, by_dim[dim])

        if len(by_dim) > 10:
            remaining = doc.add_paragraph()
            run = remaining.add_run(
                f"\n... and {len(by_dim) - 10} more dimensions not shown "
                f"({sum(len(by_dim[d]) for d in by_dim if d not in shown_dims)} members total)"
            )
            run.italic = True
            run.font.color.rgb = COLOR_TEXT_SECONDARY
            run.font.size = Pt(10)
    else:
        doc.add_heading("Members", level=1)
        notice = doc.add_paragraph()
        notice_run = notice.add_run(
            "No members were captured. Members require an EPM Automate metadata "
            "export or LCM snapshot; run a context refresh once that is configured."
        )
        notice_run.font.color.rgb = COLOR_WARNING
        notice_run.font.size = Pt(10)

    doc.add_page_break()

    # --- rules ---
    rules = _sorted_records(records, "rule")
    if rules:
        doc.add_heading(f"Business Rules ({len(rules)})", level=1)

        intro = doc.add_paragraph()
        intro.add_run(
            "Calculation scripts and business rules defined in the application."
        ).font.size = Pt(10)

        for idx, r in enumerate(rules, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)

            # Number and name
            run_num = p.add_run(f"{idx}. ")
            run_num.font.bold = True
            run_num.font.color.rgb = COLOR_PRIMARY
            run_num.font.size = Pt(10)

            run_name = p.add_run(r.name)
            run_name.font.size = Pt(10)

    doc.add_page_break()

    # --- variables ---
    variables = _sorted_records(records, "variable")
    if variables:
        doc.add_heading(f"Substitution & User Variables ({len(variables)})", level=1)

        intro = doc.add_paragraph()
        intro.add_run(
            "Variables used for dynamic calculations and member substitutions."
        ).font.size = Pt(10)

        rows = []
        for v in variables:
            data = v.data or {}
            rows.append([v.name, data.get("dimension") or "—", str(data.get("value") or "—")])
        _add_table(
            doc,
            ["Variable", "Dimension", "Value"],
            rows,
            header_color=COLOR_PRIMARY,
        )

    # Footer with generation info
    doc.add_page_break()
    footer_section = doc.sections[-1]
    footer = footer_section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Generated by EPM Wizard • {cv.label}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = COLOR_TEXT_SECONDARY

    buffer = BytesIO()
    doc.save(buffer)
    return f"{cv.label}.docx", buffer.getvalue()
